from docx import Document
from docx.shared import Pt
import io

def generate_docx(resume_json: dict) -> io.BytesIO:
    doc = Document()
    
    # Stylistic choices
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Contact Info
    contact = resume_json.get("contact", {})
    if contact.get("name"):
        name_p = doc.add_heading(contact["name"], 0)
        name_p.alignment = 1 # Center
    
    contact_details = []
    if contact.get("email"): contact_details.append(contact["email"])
    if contact.get("phone"): contact_details.append(contact["phone"])
    if contact.get("location"): contact_details.append(contact["location"])
    
    if contact_details:
        p = doc.add_paragraph(" | ".join(contact_details))
        p.alignment = 1

    # Professional Summary
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(resume_json.get("summary", ""))

    # Skills
    skills = resume_json.get("skills", [])
    if skills:
        doc.add_heading('Key Skills', level=1)
        doc.add_paragraph(", ".join(skills))

    # Experience
    doc.add_heading('Work Experience', level=1)
    for exp in resume_json.get("experience", []):
        p = doc.add_paragraph()
        run = p.add_run(f"{exp.get('title', 'N/A')} - {exp.get('company', 'N/A')}")
        run.bold = True
        p_period = doc.add_paragraph()
        run_period = p_period.add_run(exp.get("period", ""))
        run_period.italic = True
        for bullet in exp.get("bullets", []):
            doc.add_paragraph(bullet, style='List Bullet')

    # Education
    doc.add_heading('Education', level=1)
    for edu in resume_json.get("education", []):
        p = doc.add_paragraph()
        run = p.add_run(f"{edu.get('degree', 'N/A')}")
        run.bold = True
        doc.add_paragraph(f"{edu.get('school', 'N/A')} ({edu.get('year', 'N/A')})")

    # Save to buffer
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
